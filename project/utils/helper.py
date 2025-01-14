import abc
from datetime import datetime
import decimal
import logging
from collections import OrderedDict
from StringIO import StringIO
import re

from celery import shared_task

from django.core.cache import cache
from django.core.serializers.json import DjangoJSONEncoder
from django.shortcuts import HttpResponse
from django.utils import html

from rest_framework.renderers import JSONRenderer

import docx

import unicodecsv
import xlsxwriter


def HAWCtoDateString(datetime):
    """
    Helper function to ensure dates are consistent.
    """
    return datetime.strftime("%B %d %Y, %I:%M %p")

def cleanHTML(txt):
    return html.strip_entities(
                html.strip_tags(
                    txt.replace('\n', ' ')
                       .replace('\r', "")
                       .replace('<br>', "\n")
                       .replace("&nbsp;", " ")))


class HAWCDjangoJSONEncoder(DjangoJSONEncoder):
    """
    Modified to return a float instead of a string.
    """
    def default(self, o):
        if isinstance(o, decimal.Decimal):
            return float(o)
        else:
            return super(HAWCDjangoJSONEncoder, self).default(o)


class SerializerHelper(object):
    """
    HAWC helper-object for getting serialized objects and setting cache.
    Sets cache names based on django models and primary keys automatically.
    Sets a cache using the serialized object, and also a JSON object.
    """

    serializers = {}

    @classmethod
    def _get_cache_name(cls, model, id, json=True):
        name = "{}.{}.{}".format(model.__module__, model.__name__, id)
        if json: name += ".json"
        return name

    @classmethod
    def get_serialized(cls, obj, json=True, from_cache=True):
        if from_cache:
            name = cls._get_cache_name(obj.__class__, obj.id, json)
            cached = cache.get(name)
            if cached:
                logging.debug('using cache: {}'.format(name))
            else:
                cached = cls._serialize_and_cache(obj, json=json)
            return cached
        else:
            return cls._serialize(obj, json=json)

    @classmethod
    def _serialize(cls, obj, json=False):
        serializer = cls.serializers.get(obj.__class__)
        serialized = serializer(obj).data
        if json:
            serialized = JSONRenderer().render(serialized)
        return serialized

    @classmethod
    def _serialize_and_cache(cls, obj, json):
        # get expected object names
        name = cls._get_cache_name(obj.__class__, obj.id, json=False)
        json_name = cls._get_cache_name(obj.__class__, obj.id, json=True)

        # serialize data and get json-representation
        serialized = cls._serialize(obj, json=False)
        json_str = JSONRenderer().render(serialized)
        serialized = OrderedDict(serialized)  # for pickling

        logging.debug('setting cache: {}'.format(name))
        cache.set_many({name: serialized, json_name: json_str})

        if json:
            return json_str
        else:
            return serialized

    @classmethod
    def add_serializer(cls, model, serializer):
        cls.serializers[model] = serializer

    @classmethod
    def delete_caches(cls, model, ids):
        names = [cls._get_cache_name(model, id, json=False) for id in ids]
        names.extend([cls._get_cache_name(model, id, json=True) for id in ids])
        logging.debug("Removing caches: {}".format(', '.join(names)))
        cache.delete_many(names)


class FlatFileExporter(object):
    """
    Base class used to generate flat-file exports of serialized data.
    """
    def __init__(self, queryset, export_format, **kwargs):
        self.queryset = queryset
        self.export_format = export_format
        self.kwargs = kwargs

        if self.export_format == "tsv":
            self.exporter = TSVFileBuilder(**kwargs)
        elif self.export_format == "excel":
            self.exporter = ExcelFileBuilder(**kwargs)
        else:
            raise ValueError("export_format not found: {}".format(self.export_format))

    def _get_header_row(self):
        raise NotImplementedError()

    def _get_data_rows(self):
        raise NotImplementedError()

    def build_response(self):
        header_row = self._get_header_row()
        data_rows = self._get_data_rows()
        return self.exporter.generate_response(header_row, data_rows)


class FlatFile(object):
    """
    Generic file-builder object, providing an interface for generation of
    some-type of flat-file-export.

    Optional initialization argument:

        - `filename`: String filename, without extension (default: "download")
    """

    def __init__(self, filename="download", **kwargs):
        self.filename = filename
        self.kwargs = kwargs

    def generate_response(self, header_row, data_rows):
        self._setup()
        self._write_header_row(header_row)
        self._write_data_rows(data_rows)
        return self._django_response()

    def _setup(self):
        raise NotImplementedError()

    def _write_header_row(self, header_row):
        # `header_row` is a list of strings
        raise NotImplementedError()

    def _write_data_rows(self, data_rows):
        # `data_rows` is a list of lists
        raise NotImplementedError()

    def _django_response(self):
        raise NotImplementedError()


class ExcelFileBuilder(FlatFile):
    """
    Implementation of FlatFile to generate an Excel workbook with a single
    Excel worksheet. Has one header row with minor styles applied.

    Optional initialization argument:

    - `sheet_name`: String name of worksheet (default: "Sheet1")

    """

    def _setup(self):
        self.output = StringIO()
        self.wb = xlsxwriter.Workbook(self.output)
        self._add_worksheet(sheet_name=self.kwargs.get("sheet_name", "Sheet1"))

    def _add_worksheet(self, sheet_name="Sheet1"):
        """
        Create a new blank worksheet, and make sure the worksheet name is valid:
        - Make sure the name you entered does not exceed 31 characters.
        - Make sure the name does not contain any of the following characters: : \ / ? * [ or ]
        - Make sure you did not leave the name blank.
        http://stackoverflow.com/questions/451452/
        """
        sheet_name = re.sub(r'[\:\\/\?\*\[\]]+', r'-', sheet_name)[:31]
        self.ws = self.wb.add_worksheet(sheet_name)

    def _write_header_row(self, header_row):
        # set formatting and freeze panes for header-row
        header_fmt = self.wb.add_format({'bold': True})
        self.ws.freeze_panes(1, 0)
        self.ncols = len(header_row)

        # write header-rows
        for col, val in enumerate(header_row):
            self.ws.write(0, col, val, header_fmt)

    def _write_data_rows(self, data_rows):
        date_format = self.wb.add_format({'num_format': 'dd/mm/yy'})

        def write_cell(r, c, val):
            if type(val) is bool:
                return self.ws.write_boolean(r, c, val)
            elif type(val) is datetime:
                return self.ws.write_datetime(r, c, val.replace(tzinfo=None), date_format)

            try:
                val = float(val)
            except:
                pass

            return self.ws.write(r, c, val)

        r = 0
        for row in data_rows:
            r += 1
            for c, val in enumerate(row):
                write_cell(r, c, val)

        self.ws.autofilter(0, 0, r, self.ncols)

    def _django_response(self):
        fn = '{}.xlsx'.format(self.filename)
        self.wb.close()
        self.output.seek(0)
        response = HttpResponse(self.output.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="{}"'.format(fn)
        return response


class TSVFileBuilder(FlatFile):
    """
    Implementation of FlatFile to generate an tab-separated value file.
    """

    def _setup(self):
        self.output = StringIO()
        self.tsv = unicodecsv.writer(self.output, dialect='excel-tab', encoding='utf-8')

    def _write_header_row(self, header_row):
        self.tsv.writerow(header_row)

    def _write_data_rows(self, data_rows):
        self.tsv.writerows(data_rows)

    def _django_response(self):
        self.output.seek(0)
        response = HttpResponse(self.output, content_type='text/tab-separated-values')
        response['Content-Disposition'] = 'attachment; filename="{}.tsv"'.format(self.filename)
        return response


class DOCXReport(object):

    def __init__(self, fn, context):
        self.fn = fn
        self.context = context

    @shared_task
    def build_report(self):
        """
        Build DOCX report, apply context, and return Django HTTP response
        """
        self.doc = docx.Document()
        self.create_context()
        return self.django_response()

    @abc.abstractmethod
    def create_context(self):
        """
        Main-method called to generate the content in a Word report
        """
        pass

    def django_response(self):
        """
        Create an HttpResponse object with the appropriate headers.
        """
        docx_file = StringIO()
        self.doc.save(docx_file)
        docx_file.seek(0)
        response = HttpResponse(docx_file)
        response['Content-Disposition'] = 'attachment; filename={}'.format(self.fn)
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response

    def build_table(self, numRows, numCols, cells, numHeaders=1):
        """
        Helper function to build a table.

        - numRows: (int)
        - numCols: (int)
        - cells: (list) in the following format:
            [
                {"row":0, "col":0, "text":"value"},
                {"row":1, "col":0, "text":"value", "rowspan": 2},
                {"row":1, "col":0, "text":"value", "colspan": 2},
                {"row":0, "col":0, "runs":[
                    {"text": "standard text"},
                    {"text": "bolded text", "bold": True}
                ]},
            ]

        """

        tbl = self.doc.add_table(rows=numRows, cols=numCols, style='Light Shading')
        tbl.autofit = False

        for cell in cells:
            cellD = tbl.cell(cell["row"], cell["col"])
            p = cellD.paragraphs[0]

            # merge cells if needed
            rowspan = cell.get("rowspan", None)
            colspan = cell.get("colspan", None)
            if rowspan or colspan:
                rowIdx = cell["row"] + cell.get("rowspan", 1) - 1
                colIdx = cell["col"] + cell.get("colspan", 1) - 1
                cellD.merge(tbl.cell(rowIdx, colIdx))

            # add cell-shading if needed
            if "shade" in cell:
                shade_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
                    docx.oxml.ns.nsdecls('w'), cell["shade"]))
                cellD._tc.get_or_add_tcPr().append(shade_elm)

            # add content
            if "width" in cell:
                cellD.width = docx.shared.Inches(cell["width"])

            if "text" in cell:
                p.text = cell["text"]

            if "runs" in cell:
                for runD in cell["runs"]:
                    run = p.add_run(runD["text"])
                    run.bold = runD.get("bold", False)
                    run.italic = runD.get("italic", False)

        # mark rows as headers to break on pages
        if numHeaders>=1:
            for i in xrange(numHeaders):
                tblHeader = docx.oxml.parse_xml(r'<w:tblHeader {} />'.format(
                    docx.oxml.ns.nsdecls('w')))
                tbl.rows[i]._tr.get_or_add_trPr().append(tblHeader)

        return tbl
